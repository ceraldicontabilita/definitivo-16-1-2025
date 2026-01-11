"""
Employee Contracts Router - Gestione contratti dipendenti.
"""
from fastapi import APIRouter, HTTPException, Body
from fastapi.responses import FileResponse
from typing import Dict, Any, List
from datetime import datetime
import logging
import os
import uuid
import shutil
from docx import Document
import tempfile

from app.database import Database, Collections

logger = logging.getLogger(__name__)
router = APIRouter()

# Contract templates directory
CONTRACTS_DIR = "/app/uploads/contracts"
TEMPLATES_DIR = "/app/uploads/contract_templates"

# Available contract types
CONTRACT_TYPES = [
    {"id": "determinato", "name": "Contratto a Tempo Determinato", "filename": "Contratto derminato.docx"},
    {"id": "indeterminato", "name": "Contratto a Tempo Indeterminato", "filename": "Contratto indetermionato.docx"},
    {"id": "part_time_det", "name": "Contratto Part-Time Determinato", "filename": "Contratto part_time determinato.docx"},
    {"id": "part_time_ind", "name": "Contratto Part-Time Indeterminato", "filename": "Contratto part_time indeterminato.docx"},
    {"id": "informativa_152", "name": "Informativa D.Lgs. 152/1997", "filename": "INFORMATIVA AI SENSI DEL D.LGS. 152-1997.docx"},
    {"id": "informativa_privacy", "name": "Informativa Privacy", "filename": "Informativa-Privacy.docx"},
    {"id": "regolamento", "name": "Regolamento Interno Aziendale", "filename": "REGOLAMENTO INTERNO AZIENDALE.docx"},
    {"id": "richiesta_ferie", "name": "Richiesta Ferie", "filename": "RICHIESTA FERIE.docx"},
]


def ensure_dirs():
    """Create directories if they don't exist."""
    os.makedirs(CONTRACTS_DIR, exist_ok=True)
    os.makedirs(TEMPLATES_DIR, exist_ok=True)


def fill_contract_template(template_path: str, employee_data: Dict[str, Any]) -> str:
    """
    Fill contract template with employee data.
    Replaces specific text patterns with employee data.
    """
    doc = Document(template_path)
    
    # Build full name
    nome_completo = employee_data.get("nome_completo", "")
    if not nome_completo:
        nome_completo = f"{employee_data.get('cognome', '')} {employee_data.get('nome', '')}".strip()
    
    # All values to replace
    data_values = {
        "nome_completo": nome_completo or "______",
        "cognome": employee_data.get("cognome", "______"),
        "nome": employee_data.get("nome", "______"),
        "codice_fiscale": employee_data.get("codice_fiscale", "______"),
        "data_nascita": employee_data.get("data_nascita", "______"),
        "luogo_nascita": employee_data.get("luogo_nascita") or employee_data.get("comune_nascita", "______"),
        "indirizzo": employee_data.get("indirizzo", "______"),
        "mansione": employee_data.get("mansione") or employee_data.get("qualifica", "______"),
        "livello": employee_data.get("livello", "______"),
        "qualifica": employee_data.get("qualifica") or employee_data.get("mansione", "______"),
        "stipendio_orario": str(employee_data.get("stipendio_orario") or employee_data.get("salary", "______")),
        "data_inizio": employee_data.get("data_inizio") or employee_data.get("hire_date", "______"),
        "data_fine": employee_data.get("data_fine", "______"),
    }
    
    def replace_placeholders(text: str) -> str:
        """Replace ellipsis placeholders with employee data."""
        result = text
        
        # The template uses Unicode ellipsis character (…) repeated multiple times
        # We need to replace these patterns specifically
        
        # Pattern 1: "Lavoratore: ……………, nato a …………. il ……………………, residente in ………………………………… con codice fiscale ……………………………."
        if "Lavoratore:" in result and "…" in result:
            # Replace the entire line
            result = f"Lavoratore: {data_values['mansione']}, {data_values['nome_completo']}, nato a {data_values['luogo_nascita']} il {data_values['data_nascita']}, residente in {data_values['indirizzo']} con codice fiscale {data_values['codice_fiscale']}."
        
        # Pattern 2: "IL Sig. ……………………………. è assunto"
        elif "IL Sig." in result and "…" in result:
            import re
            result = re.sub(r'IL Sig\.\s*[…\.]+\s*è assunto', f"IL Sig. {data_values['nome_completo']} è assunto", result)
        
        # Pattern 3: "mansioni: ………………………… inquadrato"
        elif "mansioni:" in result and "…" in result:
            import re
            result = re.sub(r'mansioni:\s*[…\.]+\s*inquadrato', f"mansioni: {data_values['mansione']} inquadrato", result)
        
        # Pattern 4: "livello …….." or "livello …………"
        elif "livello" in result.lower() and "…" in result:
            import re
            result = re.sub(r'livello\s*[…\.]+', f"livello {data_values['livello']}", result, flags=re.IGNORECASE)
        
        # Pattern 5: "qualifica …………" 
        elif "qualifica" in result.lower() and "…" in result:
            import re
            result = re.sub(r'qualifica\s*[…\.]+', f"qualifica {data_values['qualifica']}", result, flags=re.IGNORECASE)
        
        # Pattern 6: "euro ………………… ora" (stipendio)
        elif "euro" in result.lower() and "ora" in result.lower() and "…" in result:
            import re
            result = re.sub(r'euro\s*[…\.]+\s*ora', f"euro {data_values['stipendio_orario']} ora", result, flags=re.IGNORECASE)
        
        # Pattern 7: "decorre dal ………… al …………"
        elif "decorre dal" in result.lower() and "…" in result:
            import re
            result = re.sub(r'decorre dal\s*[…\.]+\s*al\s*[…\.]+', f"decorre dal {data_values['data_inizio']} al {data_values['data_fine']}", result, flags=re.IGNORECASE)
        
        # Generic fallback: replace any remaining ellipsis sequences
        elif "…" in result:
            import re
            # Replace sequences of 10+ ellipsis with longer values
            result = re.sub(r'[…]{10,}', data_values['nome_completo'], result)
            # Replace sequences of 6-9 ellipsis with medium values
            result = re.sub(r'[…]{6,9}', data_values['mansione'], result)
            # Replace sequences of 3-5 ellipsis with shorter values
            result = re.sub(r'[…]{3,5}', "______", result)
        
        return result
    
    # Process all paragraphs
    for para in doc.paragraphs:
        if "…" in para.text or "…" in para.text:
            # Process the entire paragraph text
            new_text = replace_placeholders(para.text)
            if new_text != para.text:
                # Clear runs and set new text
                if para.runs:
                    # Keep first run's formatting
                    first_run = para.runs[0]
                    for run in para.runs[1:]:
                        run.text = ""
                    first_run.text = new_text
                else:
                    para.text = new_text
    
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if "…" in para.text:
                        new_text = replace_placeholders(para.text)
                        if new_text != para.text:
                            if para.runs:
                                first_run = para.runs[0]
                                for run in para.runs[1:]:
                                    run.text = ""
                                first_run.text = new_text
                            else:
                                para.text = new_text
    
    # Save to temp file
    output_path = tempfile.mktemp(suffix=".docx")
    doc.save(output_path)
    
    return output_path


@router.get("/types")
async def get_contract_types() -> List[Dict[str, str]]:
    """Get available contract types."""
    return CONTRACT_TYPES


@router.get("/templates")
async def list_templates() -> List[Dict[str, Any]]:
    """List available contract templates."""
    ensure_dirs()
    
    templates = []
    for ct in CONTRACT_TYPES:
        template_path = os.path.join(TEMPLATES_DIR, ct["filename"])
        exists = os.path.exists(template_path)
        templates.append({
            "id": ct["id"],
            "name": ct["name"],
            "filename": ct["filename"],
            "available": exists
        })
    
    return templates


@router.post("/generate/{employee_id}")
async def generate_contract(employee_id: str, data: Dict[str, Any] = Body(...)) -> Dict[str, Any]:
    """
    Generate a contract for an employee.
    
    Request body:
    {
        "contract_type": "determinato",
        "additional_data": {
            "livello": "5",
            "stipendio_orario": "8.50",
            "qualifica": "Barista"
        }
    }
    """
    ensure_dirs()
    
    contract_type = data.get("contract_type") or data.get("contract_type_id")
    additional_data = data.get("additional_data", {})
    
    # Find contract type
    ct = next((c for c in CONTRACT_TYPES if c["id"] == contract_type), None)
    if not ct:
        raise HTTPException(status_code=400, detail=f"Tipo contratto non valido: {contract_type}")
    
    # Get employee
    db = Database.get_db()
    employee = await db[Collections.EMPLOYEES].find_one({"id": employee_id}, {"_id": 0})
    
    if not employee:
        raise HTTPException(status_code=404, detail="Dipendente non trovato")
    
    # Check template exists
    template_path = os.path.join(TEMPLATES_DIR, ct["filename"])
    if not os.path.exists(template_path):
        # Try fallback path
        fallback_path = f"/tmp/documenti contratti dipendente/{ct['filename']}"
        if os.path.exists(fallback_path):
            template_path = fallback_path
        else:
            raise HTTPException(status_code=404, detail=f"Template non trovato: {ct['filename']}")
    
    # Merge employee data with additional data
    employee_data = {**employee, **additional_data}
    
    # Format date if present
    if employee_data.get("data_nascita"):
        try:
            dt = datetime.fromisoformat(str(employee_data["data_nascita"]).replace("Z", ""))
            employee_data["data_nascita"] = dt.strftime("%d/%m/%Y")
        except (ValueError, TypeError):
            pass
    
    try:
        # Generate filled contract
        output_path = fill_contract_template(template_path, employee_data)
        
        # Create final filename
        safe_name = employee_data.get("nome_completo", "dipendente").replace(" ", "_")
        final_filename = f"{ct['id']}_{safe_name}_{datetime.now().strftime('%Y%m%d')}.docx"
        final_path = os.path.join(CONTRACTS_DIR, final_filename)
        
        # Move to contracts dir
        shutil.move(output_path, final_path)
        
        # Record contract generation
        contract_record = {
            "id": str(uuid.uuid4()),
            "employee_id": employee_id,
            "employee_name": employee_data.get("nome_completo"),
            "contract_type": contract_type,
            "contract_name": ct["name"],
            "filename": final_filename,
            "filepath": final_path,
            "generated_at": datetime.utcnow().isoformat(),
            "additional_data": additional_data
        }
        
        await db["employee_contracts"].insert_one(contract_record)
        
        return {
            "success": True,
            "message": f"Contratto generato per {employee_data.get('nome_completo')}",
            "contract": {
                "id": contract_record["id"],
                "filename": final_filename,
                "download_url": f"/api/contracts/download/{contract_record['id']}"
            }
        }
        
    except Exception as e:
        logger.error(f"Error generating contract: {e}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Errore generazione contratto: {str(e)}")


@router.get("/download/{contract_id}")
async def download_contract(contract_id: str):
    """Download a generated contract."""
    db = Database.get_db()
    contract = await db["employee_contracts"].find_one({"id": contract_id}, {"_id": 0})
    
    if not contract:
        raise HTTPException(status_code=404, detail="Contratto non trovato")
    
    filepath = contract.get("filepath")
    if not filepath or not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail="File contratto non trovato")
    
    return FileResponse(
        filepath,
        filename=contract.get("filename"),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@router.get("/employee/{employee_id}")
async def get_employee_contracts(employee_id: str) -> List[Dict[str, Any]]:
    """Get all contracts for an employee."""
    db = Database.get_db()
    contracts = await db["employee_contracts"].find(
        {"employee_id": employee_id},
        {"_id": 0}
    ).sort("generated_at", -1).to_list(100)
    
    return contracts


@router.delete("/{contract_id}")
async def delete_contract(contract_id: str) -> Dict[str, Any]:
    """Delete a generated contract."""
    db = Database.get_db()
    contract = await db["employee_contracts"].find_one({"id": contract_id}, {"_id": 0})
    
    if not contract:
        raise HTTPException(status_code=404, detail="Contratto non trovato")
    
    # Delete file
    filepath = contract.get("filepath")
    if filepath and os.path.exists(filepath):
        os.remove(filepath)
    
    # Delete record
    await db["employee_contracts"].delete_one({"id": contract_id})
    
    return {"success": True, "message": "Contratto eliminato"}
